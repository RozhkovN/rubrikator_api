"""
FastAPI приложение для классификации жалоб.

Запуск:
    uvicorn api.main:app --reload --host 0.0.0.0 --port 8800
"""

import os
import sys
from pathlib import Path

# Добавляем корневую директорию в путь
root_dir = Path(__file__).parent.parent
sys.path.insert(0, str(root_dir))

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import Optional, List, Dict
import logging
import io

# Для работы с документами
import subprocess
import tempfile

# python-docx для .docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx не установлен. Загрузка .docx файлов недоступна.")

# pypdf для PDF
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("pypdf не установлен. Загрузка PDF файлов недоступна.")

# olefile для .doc
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False
    logging.warning("olefile не установлен.")

from src.model import ComplaintClassifier

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Создание FastAPI приложения
app = FastAPI(
    title="Рубрикатор Жалоб API",
    description="API для автоматической классификации жалоб по 20 рубрикаторам",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Глобальная переменная для классификатора
classifier: Optional[ComplaintClassifier] = None


# Модели данных
class ComplaintRequest(BaseModel):
    """Запрос на классификацию жалобы"""
    text: str = Field(..., description="Текст жалобы", min_length=10)
    top_k: int = Field(1, description="Количество топ результатов", ge=1, le=5)
    
    class Config:
        json_schema_extra = {
            "example": {
                "text": "Банк заблокировал мою карту без предупреждения согласно 161-ФЗ",
                "top_k": 1
            }
        }


class PredictionItem(BaseModel):
    """Один результат классификации"""
    rubric_id: int
    rubric_name: str
    short_name: str
    response_template: str
    confidence: float


class ComplaintResponse(BaseModel):
    """Ответ с результатами классификации"""
    text: str
    best_match: PredictionItem
    all_predictions: Optional[List[PredictionItem]] = None


class HealthResponse(BaseModel):
    """Ответ health check"""
    status: str
    model_loaded: bool


class TrainRequest(BaseModel):
    """Запрос на обучение модели"""
    model_name: str = Field(
        "paraphrase-multilingual-mpnet-base-v2",
        description="Название модели Sentence Transformers"
    )
    use_keywords: bool = Field(True, description="Использовать ли анализ ключевых слов")
    keyword_weight: float = Field(0.3, description="Вес ключевых слов (0-1)", ge=0, le=1)
    
    class Config:
        json_schema_extra = {
            "example": {
                "model_name": "paraphrase-multilingual-mpnet-base-v2",
                "use_keywords": True,
                "keyword_weight": 0.3
            }
        }


class TrainResponse(BaseModel):
    """Ответ на запрос обучения"""
    status: str
    message: str
    model_path: str


# События жизненного цикла
@app.on_event("startup")
async def startup_event():
    """Загрузка модели при старте приложения"""
    global classifier
    
    logger.info("🚀 Запуск API сервера...")
    
    try:
        model_path = root_dir / "models" / "classifier.pkl"
        
        if not model_path.exists():
            logger.warning(f"⚠️  Модель не найдена: {model_path}")
            logger.warning("Используйте POST /train для обучения модели")
            classifier = None
        else:
            classifier = ComplaintClassifier()
            classifier.load(str(model_path))
            logger.info("✅ Классификатор загружен успешно")
        
    except Exception as e:
        logger.error(f"❌ Ошибка загрузки модели: {e}")
        logger.warning("API запущен без модели. Используйте POST /train для обучения")
        classifier = None


@app.on_event("shutdown")
async def shutdown_event():
    """Очистка ресурсов при остановке"""
    logger.info("🛑 Остановка API сервера...")


# Эндпоинты
@app.get("/", tags=["General"])
async def root():
    """Корневой эндпоинт"""
    return {
        "message": "Рубрикатор Жалоб API",
        "version": "1.0.0",
        "docs": "/docs"
    }


@app.get("/health", response_model=HealthResponse, tags=["General"])
async def health():
    """Проверка состояния API"""
    return {
        "status": "ok",
        "model_loaded": classifier is not None
    }


@app.post("/classify", response_model=ComplaintResponse, tags=["Classification"])
async def classify_complaint(request: ComplaintRequest):
    """
    Классификация жалобы.
    
    Принимает текст жалобы и возвращает наиболее подходящий рубрикатор.
    
    - **text**: Текст жалобы (минимум 10 символов)
    - **top_k**: Количество топ результатов (по умолчанию 1)
    
    Возвращает:
    - **best_match**: Лучшее совпадение с рубрикатором
    - **all_predictions**: Все топ-k результатов (если top_k > 1)
    """
    if classifier is None:
        raise HTTPException(
            status_code=503,
            detail="Классификатор не загружен. Сервер не готов к работе."
        )
    
    try:
        # Классифицируем жалобу
        result = classifier.predict(
            text=request.text,
            top_k=request.top_k,
            return_scores=False
        )
        
        # Формируем ответ
        predictions = []
        for pred in result['predictions']:
            predictions.append(PredictionItem(
                rubric_id=pred['rubric_id'],
                rubric_name=pred['rubric_name'],
                short_name=pred.get('short_name', ''),
                response_template=pred.get('response_template', ''),
                confidence=round(pred['confidence'], 4)
            ))
        
        response = ComplaintResponse(
            text=request.text,
            best_match=predictions[0],
            all_predictions=predictions if request.top_k > 1 else None
        )
        
        logger.info(f"✅ Классифицировано: {predictions[0].short_name} ({predictions[0].confidence:.2%})")
        
        return response
        
    except Exception as e:
        logger.error(f"❌ Ошибка классификации: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка при классификации: {str(e)}"
        )


@app.post("/classify/batch", tags=["Classification"])
async def classify_batch(complaints: List[str]):
    """
    Пакетная классификация жалоб.
    
    Принимает список текстов жалоб и возвращает результаты для каждой.
    """
    if classifier is None:
        raise HTTPException(
            status_code=503,
            detail="Классификатор не загружен"
        )
    
    if len(complaints) > 100:
        raise HTTPException(
            status_code=400,
            detail="Максимум 100 жалоб за один запрос"
        )
    
    try:
        results = classifier.predict_batch(complaints, top_k=1)
        
        responses = []
        for result in results:
            pred = result['best_match']
            responses.append({
                "text": result['text'],
                "rubric_id": pred['rubric_id'],
                "rubric_name": pred['rubric_name'],
                "short_name": pred.get('short_name', ''),
                "response_template": pred.get('response_template', ''),
                "confidence": round(pred['confidence'], 4)
            })
        
        return {"results": responses, "count": len(responses)}
        
    except Exception as e:
        logger.error(f"❌ Ошибка пакетной классификации: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка при классификации: {str(e)}"
        )


@app.post("/train", response_model=TrainResponse, tags=["Model Management"])
async def train_model(request: TrainRequest):
    """
    Обучение (подготовка) модели классификатора.
    
    Создает векторные представления рубрикаторов и сохраняет модель.
    После обучения модель автоматически загружается в memory.
    
    - **model_name**: Название модели Sentence Transformers
    - **use_keywords**: Использовать ли анализ ключевых слов
    - **keyword_weight**: Вес ключевых слов (0-1)
    
    Примечание: Процесс может занять несколько минут при первом запуске
    (загрузка модели из интернета).
    """
    global classifier
    
    try:
        logger.info("🚀 Начинаем обучение модели...")
        logger.info(f"   Модель: {request.model_name}")
        logger.info(f"   Ключевые слова: {request.use_keywords}")
        logger.info(f"   Вес ключевых слов: {request.keyword_weight}")
        
        # Создаем новый классификатор с заданными параметрами
        new_classifier = ComplaintClassifier(
            model_name=request.model_name,
            use_keywords=request.use_keywords,
            keyword_weight=request.keyword_weight
        )
        
        # Путь для сохранения модели
        model_path = root_dir / "models" / "classifier.pkl"
        
        # Обучаем (создаем эмбеддинги)
        new_classifier.train(save_path=str(model_path))
        
        # Заменяем глобальный классификатор на новый
        classifier = new_classifier
        
        logger.info("✅ Обучение завершено успешно")
        
        return TrainResponse(
            status="success",
            message="Модель успешно обучена и загружена",
            model_path=str(model_path)
        )
        
    except Exception as e:
        logger.error(f"❌ Ошибка обучения модели: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка при обучении модели: {str(e)}"
        )


@app.get("/model/info", tags=["Model Management"])
async def get_model_info():
    """
    Получение информации о текущей загруженной модели.
    
    Возвращает параметры и статус модели.
    """
    if classifier is None:
        return {
            "loaded": False,
            "message": "Модель не загружена"
        }
    
    return {
        "loaded": True,
        "model_name": classifier.model_name,
        "use_keywords": classifier.use_keywords,
        "keyword_weight": classifier.keyword_weight,
        "semantic_weight": classifier.semantic_weight,
        "rubrics_count": len(classifier.rubrics)
    }


# ===================== ЗАГРУЗКА ДОКУМЕНТОВ =====================

SUPPORTED_EXTENSIONS = ['.docx', '.doc', '.pdf']


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Извлекает текст из содержимого .docx файла.
    
    Args:
        file_content: Бинарное содержимое .docx файла
    
    Returns:
        Извлеченный текст
    """
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="python-docx не установлен. Установите: pip install python-docx"
        )
    
    try:
        doc = Document(io.BytesIO(file_content))
        
        # Извлекаем текст из всех параграфов
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        
        full_text = "\n".join(paragraphs)
        return full_text
        
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Ошибка чтения .docx файла: {str(e)}"
        )


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Извлекает текст из содержимого PDF файла.
    
    Args:
        file_content: Бинарное содержимое PDF файла
    
    Returns:
        Извлеченный текст
    """
    if not PDF_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="pypdf не установлен. Установите: pip install pypdf"
        )
    
    try:
        reader = PdfReader(io.BytesIO(file_content))
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())
        
        full_text = "\n".join(text_parts)
        return full_text
        
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Ошибка чтения PDF файла: {str(e)}"
        )


def extract_text_from_doc(file_content: bytes) -> str:
    """
    Извлекает текст из содержимого .doc файла (старый формат Word).
    
    Использует antiword (если доступен) или olefile для базового извлечения.
    
    Args:
        file_content: Бинарное содержимое .doc файла
    
    Returns:
        Извлеченный текст
    """
    # Сначала пробуем antiword (лучший результат)
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        try:
            result = subprocess.run(
                ['antiword', '-w', '0', tmp_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (subprocess.SubprocessError, FileNotFoundError):
            pass  # antiword не доступен, пробуем olefile
        finally:
            import os
            os.unlink(tmp_path)
    except Exception:
        pass
    
    # Пробуем olefile для базового извлечения
    if OLEFILE_AVAILABLE:
        try:
            ole = olefile.OleFileIO(io.BytesIO(file_content))
            
            # Пробуем извлечь текст из WordDocument stream
            if ole.exists('WordDocument'):
                # Базовое извлечение текста (может быть неполным)
                text_parts = []
                
                # Ищем текстовые потоки
                for stream in ole.listdir():
                    stream_name = '/'.join(stream)
                    if 'text' in stream_name.lower() or stream == ['WordDocument']:
                        try:
                            data = ole.openstream(stream).read()
                            # Пытаемся декодировать как текст
                            try:
                                text = data.decode('utf-16-le', errors='ignore')
                                # Оставляем только печатные символы
                                clean_text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t ')
                                if len(clean_text) > 20:
                                    text_parts.append(clean_text)
                            except:
                                pass
                        except:
                            pass
                
                ole.close()
                
                if text_parts:
                    return "\n".join(text_parts)
            
            ole.close()
        except Exception as e:
            logger.warning(f"olefile не смог прочитать .doc: {e}")
    
    raise HTTPException(
        status_code=400,
        detail="Не удалось извлечь текст из .doc файла. Рекомендуем конвертировать в .docx или .pdf"
    )


def extract_text_from_file(filename: str, content: bytes) -> str:
    """
    Универсальная функция извлечения текста из файла.
    
    Args:
        filename: Имя файла
        content: Бинарное содержимое файла
    
    Returns:
        Извлеченный текст
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    
    if ext == 'docx':
        return extract_text_from_docx(content)
    elif ext == 'doc':
        return extract_text_from_doc(content)
    elif ext == 'pdf':
        return extract_text_from_pdf(content)
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый формат файла: .{ext}. Поддерживаются: .docx, .doc, .pdf"
        )


@app.post("/classify/file", response_model=ComplaintResponse, tags=["Classification"])
async def classify_from_file(
    file: UploadFile = File(..., description="Документ с текстом жалобы (.docx, .doc, .pdf)"),
    top_k: int = 1
):
    """
    Классификация жалобы из файла документа.
    
    Загружает файл, извлекает текст и классифицирует жалобу.
    
    Поддерживаемые форматы:
    - **.docx** - Microsoft Word (современный формат)
    - **.doc** - Microsoft Word (старый формат)
    - **.pdf** - PDF документ
    
    - **top_k**: Количество топ результатов (по умолчанию 1)
    """
    if classifier is None:
        raise HTTPException(
            status_code=503,
            detail="Классификатор не загружен. Сервер не готов к работе."
        )
    
    # Проверяем расширение файла
    ext = file.filename.lower().rsplit('.', 1)[-1] if '.' in file.filename else ''
    if f'.{ext}' not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый формат. Поддерживаются: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    
    # Проверяем размер файла (максимум 10 МБ)
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(
            status_code=400,
            detail="Размер файла превышает 10 МБ"
        )
    
    # Извлекаем текст
    text = extract_text_from_file(file.filename, content)
    
    if len(text.strip()) < 10:
        raise HTTPException(
            status_code=400,
            detail="Извлеченный текст слишком короткий (менее 10 символов)"
        )
    
    try:
        # Классифицируем
        result = classifier.predict(
            text=text,
            top_k=top_k,
            return_scores=False
        )
        
        # Формируем ответ
        predictions = []
        for pred in result['predictions']:
            predictions.append(PredictionItem(
                rubric_id=pred['rubric_id'],
                rubric_name=pred['rubric_name'],
                short_name=pred.get('short_name', ''),
                response_template=pred.get('response_template', ''),
                confidence=round(pred['confidence'], 4)
            ))
        
        response = ComplaintResponse(
            text=text,
            best_match=predictions[0],
            all_predictions=predictions if top_k > 1 else None
        )
        
        logger.info(f"✅ Файл {file.filename} классифицирован: {predictions[0].short_name}")
        
        return response
        
    except Exception as e:
        logger.error(f"❌ Ошибка классификации файла: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка при классификации: {str(e)}"
        )


class FileClassificationResult(BaseModel):
    """Результат классификации одного файла"""
    filename: str
    text: str
    best_match: Optional[PredictionItem] = None
    all_predictions: Optional[List[PredictionItem]] = None
    error: Optional[str] = None


class FilesClassificationResponse(BaseModel):
    """Ответ пакетной классификации файлов"""
    results: List[FileClassificationResult]
    total: int
    success: int
    failed: int


@app.post("/classify/files", response_model=FilesClassificationResponse, tags=["Classification"])
async def classify_from_files(
    files: List[UploadFile] = File(..., description="Список документов (.docx, .doc, .pdf)"),
    top_k: int = 1
):
    """
    Пакетная классификация жалоб из нескольких файлов.
    
    Загружает до 10 файлов, извлекает текст и классифицирует каждую жалобу.
    
    Поддерживаемые форматы:
    - **.docx** - Microsoft Word (современный формат)
    - **.doc** - Microsoft Word (старый формат)
    - **.pdf** - PDF документ
    
    - **files**: Список файлов, максимум 10
    - **top_k**: Количество топ результатов для каждого файла (1-5)
    """
    if classifier is None:
        raise HTTPException(
            status_code=503,
            detail="Классификатор не загружен. Сервер не готов к работе."
        )
    
    # Проверяем top_k
    if top_k < 1 or top_k > 5:
        raise HTTPException(
            status_code=400,
            detail="top_k должен быть от 1 до 5"
        )
    
    # Проверяем количество файлов
    if len(files) > 10:
        raise HTTPException(
            status_code=400,
            detail="Максимум 10 файлов за один запрос"
        )
    
    results = []
    success_count = 0
    failed_count = 0
    
    for file in files:
        try:
            # Проверяем расширение
            ext = file.filename.lower().rsplit('.', 1)[-1] if '.' in file.filename else ''
            if f'.{ext}' not in SUPPORTED_EXTENSIONS:
                results.append(FileClassificationResult(
                    filename=file.filename,
                    text="",
                    best_match=None,
                    all_predictions=None,
                    error=f"Неподдерживаемый формат. Поддерживаются: {', '.join(SUPPORTED_EXTENSIONS)}"
                ))
                failed_count += 1
                continue
            
            # Читаем содержимое
            content = await file.read()
            
            # Проверяем размер
            if len(content) > 10 * 1024 * 1024:
                results.append(FileClassificationResult(
                    filename=file.filename,
                    text="",
                    best_match=None,
                    all_predictions=None,
                    error="Размер файла превышает 10 МБ"
                ))
                failed_count += 1
                continue
            
            # Извлекаем текст
            text = extract_text_from_file(file.filename, content)
            
            if len(text.strip()) < 10:
                results.append(FileClassificationResult(
                    filename=file.filename,
                    text=text,
                    best_match=None,
                    all_predictions=None,
                    error="Текст слишком короткий"
                ))
                failed_count += 1
                continue
            
            # Классифицируем
            result = classifier.predict(text=text, top_k=top_k, return_scores=False)
            
            # Формируем предсказания
            predictions = []
            for pred in result['predictions']:
                predictions.append(PredictionItem(
                    rubric_id=pred['rubric_id'],
                    rubric_name=pred['rubric_name'],
                    short_name=pred.get('short_name', ''),
                    response_template=pred.get('response_template', ''),
                    confidence=round(pred['confidence'], 4)
                ))
            
            results.append(FileClassificationResult(
                filename=file.filename,
                text=text,
                best_match=predictions[0],
                all_predictions=predictions if top_k > 1 else None
            ))
            success_count += 1
            logger.info(f"✅ Файл {file.filename}: {predictions[0].short_name}")
            
        except Exception as e:
            results.append(FileClassificationResult(
                filename=file.filename,
                text="",
                best_match=None,
                all_predictions=None,
                error=str(e)
            ))
            failed_count += 1
            logger.error(f"❌ Ошибка файла {file.filename}: {e}")
    
    return FilesClassificationResponse(
        results=results,
        total=len(files),
        success=success_count,
        failed=failed_count
    )


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8800))
    uvicorn.run(app, host="0.0.0.0", port=port)
